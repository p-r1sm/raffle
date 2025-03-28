import os
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Mm
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont
import io

def use_external_logo(logo_path='resources/logo.png', src_image=None):
    """Use an external logo image or create one if not provided."""
    # If the source image is provided, save it to the logo path
    if src_image and os.path.exists(src_image):
        # Ensure directory exists
        os.makedirs(os.path.dirname(logo_path), exist_ok=True)
        
        # Copy the image
        try:
            from shutil import copyfile
            copyfile(src_image, logo_path)
            return logo_path
        except Exception as e:
            print(f"Error copying logo: {e}")
    
    # Create a logo if one doesn't exist
    return create_circular_logo(logo_path)

def create_circular_logo(output_path='resources/logo.png'):
    """Create a circular logo image similar to the one in the reference."""
    # Check if logo already exists
    if os.path.exists(output_path):
        return output_path
    
    # Ensure directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    # Create a circular logo with a silhouette
    size = (300, 300)
    img = Image.new('RGBA', size, (0, 0, 0, 0))
    draw = ImageDraw.Draw(img)
    
    # Draw golden circle
    circle_color = (192, 155, 85, 255)  # Golden color
    draw.ellipse((0, 0, 300, 300), fill=circle_color)
    
    # Draw white silhouette (simplified)
    white = (255, 255, 255, 255)
    # Draw a meditation pose silhouette (simplified)
    # Head
    draw.ellipse((125, 60, 175, 110), fill=white)
    # Shoulders
    draw.ellipse((110, 110, 130, 130), fill=white)
    draw.ellipse((170, 110, 190, 130), fill=white)
    # Body
    draw.rectangle((135, 110, 165, 180), fill=white)
    # Arms
    draw.polygon([(110, 130), (70, 160), (85, 180), (130, 140)], fill=white)
    draw.polygon([(190, 130), (230, 160), (215, 180), (170, 140)], fill=white)
    # Legs (folded position)
    draw.ellipse((110, 180, 190, 220), fill=white)
    
    img.save(output_path)
    return output_path

def add_border_to_table(table, color="C09B55", size=12):
    """Add a border to a table."""
    tbl = table._tbl
    for cell in table._cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        
        # Make border 3 times thicker than lines between fields
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), str(size))  # Border width in 1/8 points
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), color)
            tcBorders.append(border)
        
        tcPr.append(tcBorders)
        
        # Add background color (very light cream/beige to match image)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'FFF9E7')  # Light cream color
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

def add_horizontal_line(paragraph, color="C09B55"):
    """Add a horizontal line to a paragraph."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    
    # Horizontal line thickness (1/3 of the border thickness)
    border = OxmlElement('w:bottom')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), '4')  # Border width 
    border.set(qn('w:space'), '1')
    border.set(qn('w:color'), color)
    
    pBdr = OxmlElement('w:pBdr')
    pBdr.append(border)
    pPr.append(pBdr)

def create_card_image(data, logo_path, output_path, card_width_px=800, card_height_px=400):
    """Create a card image with the given data."""
    # Create a new image with a white background
    img = Image.new('RGB', (card_width_px, card_height_px), color='white')
    draw = ImageDraw.Draw(img)
    
    # Load fonts
    try:
        # Try to use a custom font if available
        header_font = ImageFont.truetype("resources/arial_bold.ttf", 30)
        data_font = ImageFont.truetype("resources/arial.ttf", 24)
        label_font = ImageFont.truetype("resources/arial_bold.ttf", 20)
    except IOError:
        # Fallback to default font
        header_font = ImageFont.load_default()
        data_font = ImageFont.load_default()
        label_font = ImageFont.load_default()
    
    # Define colors
    gold_color = (192, 155, 85)  # RGB for gold
    text_color = (0, 0, 0)  # Black
    
    # Draw ornate golden border
    border_width = 20
    border_color = gold_color
    draw.rectangle([0, 0, card_width_px-1, card_height_px-1], 
                   outline=border_color, 
                   width=border_width)
    
    # Draw decorative corner elements
    corner_size = 50
    corner_color = gold_color
    # Top-left
    draw.line([(0, 0), (corner_size, 0)], fill=corner_color, width=border_width//2)
    draw.line([(0, 0), (0, corner_size)], fill=corner_color, width=border_width//2)
    # Top-right
    draw.line([(card_width_px-1, 0), (card_width_px-1-corner_size, 0)], fill=corner_color, width=border_width//2)
    draw.line([(card_width_px-1, 0), (card_width_px-1, corner_size)], fill=corner_color, width=border_width//2)
    # Bottom-left
    draw.line([(0, card_height_px-1), (corner_size, card_height_px-1)], fill=corner_color, width=border_width//2)
    draw.line([(0, card_height_px-1), (0, card_height_px-1-corner_size)], fill=corner_color, width=border_width//2)
    # Bottom-right
    draw.line([(card_width_px-1, card_height_px-1), (card_width_px-1-corner_size, card_height_px-1)], fill=corner_color, width=border_width//2)
    draw.line([(card_width_px-1, card_height_px-1), (card_width_px-1, card_height_px-1-corner_size)], fill=corner_color, width=border_width//2)
    
    # Left side (text content)
    left_width = int(card_width_px * 0.65)
    
    # Vertical positions for different fields
    y_start = 80
    line_height = 50
    
    # Fields to display
    fields = [
        ("LAABHARTHI NAME", data['LAABHARTHI_NAME']),
        ("CONTACT NUMBER", data['CONTACT_NUMBER']),
        ("ARPIT GROUP", data['ARPIT_GROUP']),
        ("AREA", data['AREA'])
    ]
    
    # Draw fields
    for i, (label, value) in enumerate(fields):
        # Draw label
        draw.text((40, y_start + i*line_height), 
                  label, 
                  font=label_font, 
                  fill=gold_color)
        
        # Draw value
        draw.text((40, y_start + i*line_height + 35), 
                  value, 
                  font=data_font, 
                  fill=text_color)
        
        # Draw horizontal line
        line_y = y_start + (i+1)*line_height + 20
        draw.line([(40, line_y), (left_width-40, line_y)], 
                  fill=gold_color, 
                  width=2)
    
    # Right side (logo and amount)
    right_start_x = left_width
    
    # Add logo
    if os.path.exists(logo_path):
        logo = Image.open(logo_path)
        # Resize logo to fit
        logo_size = 250
        logo = logo.resize((logo_size, logo_size), Image.LANCZOS)
        
        # Calculate logo position
        logo_x = right_start_x + (card_width_px - right_start_x - logo_size) // 2
        logo_y = 100
        
        # Paste logo
        img.paste(logo, (logo_x, logo_y), logo if logo.mode == 'RGBA' else None)
    
    # Add amount text
    amount_text = "Amount Rs. 1000/-"
    amount_bbox = draw.textbbox((0, 0), amount_text, font=header_font)
    amount_width = amount_bbox[2] - amount_bbox[0]
    draw.text((right_start_x + (card_width_px - right_start_x - amount_width) // 2, 300), 
              amount_text, 
              font=header_font, 
              fill=gold_color)
    
    # Ensure output directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    # Save the image
    img.save(output_path)
    return output_path

def generate_cards(csv_file, output_file, rows=4, cols=2, logo_path=None):
    """Generate a Word document with card images."""
    # Generate or use existing logo
    if logo_path and os.path.exists(logo_path):
        final_logo_path = use_external_logo(src_image=logo_path)
    else:
        final_logo_path = create_circular_logo()
    
    # Read CSV file
    df = pd.read_csv(csv_file)
    
    # Convert all columns to string
    for col in df.columns:
        df[col] = df[col].astype(str)
    
    # Create document
    doc = Document()
    
    # Set page size to A4 and minimized margins
    section = doc.sections[0]
    section.page_height = Cm(29.7)  # A4 height
    section.page_width = Cm(21.0)   # A4 width
    section.left_margin = Cm(0.8)   # Further reduced margins
    section.right_margin = Cm(0.8)
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    
    # Calculate available space on the page (in points)
    available_width = Pt(section.page_width.pt - section.left_margin.pt - section.right_margin.pt)
    available_height = Pt(section.page_height.pt - section.top_margin.pt - section.bottom_margin.pt)

    # Create a temporary directory for card images
    temp_dir = os.path.join(os.path.dirname(output_file), 'temp_card_images')
    os.makedirs(temp_dir, exist_ok=True)
    
    # Generate card images
    card_images = []
    for i, record in enumerate(df.to_dict('records')):
        # Generate unique filename for each card image
        card_image_path = os.path.join(temp_dir, f'card_{i}.jpg')
        create_card_image(record, final_logo_path, card_image_path)
        card_images.append(card_image_path)
    
    # Process card images into the document
    for i in range(0, len(card_images), rows * cols):
        # Add a new page if not the first page
        if i > 0:
            doc.add_page_break()
        
        # Create a table to hold the images
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        table.autofit = False
        
        # Fill the table with images
        for r in range(rows):
            for c in range(cols):
                # Calculate the image index
                img_index = i + r * cols + c
                
                # Break if we've run out of images
                if img_index >= len(card_images):
                    break
                
                # Get the cell and add the image
                cell = table.cell(r, c)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add image to the cell
                cell.paragraphs[0].add_run().add_picture(card_images[img_index], width=Inches(2.5))
    
    # Save the document
    doc.save(output_file)
    
    # Clean up temporary image files
    for img_path in card_images:
        os.remove(img_path)
    os.rmdir(temp_dir)
    
    print(f"Cards generated successfully and saved to {output_file}")

if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description='Generate cards from CSV data')
    parser.add_argument('--csv', type=str, default='data/sample_data.csv', help='Path to CSV file')
    parser.add_argument('--output', type=str, default='output_cards.docx', help='Output Word document path')
    parser.add_argument('--rows', type=int, default=4, help='Number of rows per page')
    parser.add_argument('--cols', type=int, default=2, help='Number of columns per page')
    parser.add_argument('--logo', type=str, help='Path to custom logo image')
    
    args = parser.parse_args()
    
    # Ensure CSV file exists
    if not os.path.exists(args.csv):
        csv_path = os.path.join(os.path.dirname(__file__), args.csv)
        if not os.path.exists(csv_path):
            print(f"Error: CSV file not found at {args.csv} or {csv_path}")
            exit(1)
        args.csv = csv_path
        
    # Generate output path if it's relative
    if not os.path.isabs(args.output):
        args.output = os.path.join(os.path.dirname(__file__), args.output)
        
    # Create output directory if it doesn't exist
    output_dir = os.path.dirname(args.output)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    generate_cards(args.csv, args.output, args.rows, args.cols, args.logo) 