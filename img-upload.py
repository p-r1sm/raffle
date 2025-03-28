import os
import os.path
import pandas as pd
from PIL import Image
from docx import Document
from docx.shared import Inches, Cm
from docx.shared import Pt
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml

def get_image_files_with_ids(folder_path):
    """
    Retrieve all image files from the specified folder and match with IDs if excel provided.
    
    :param folder_path: Path to the folder containing images
    :param excel_path: Optional path to Excel file containing IDs
    :return: List of tuples (file_path, id) or just file paths if no Excel
    """
    image_files = []
    files = sorted(os.listdir(folder_path))
    sorted_files = sorted(files, key=lambda x: int(x.split('.')[0]))
    
    # Get all valid image files
    for filename in sorted_files:
        print(filename)
        full_path = os.path.join(folder_path, filename)
        try:
            # Use Pillow to check if it's a valid image
            with Image.open(full_path) as img:
                img.verify()
                image_files.append(full_path)
        except (IOError, SyntaxError):
            # Not a valid image file
            continue
    
    return image_files

def create_image_document(image_files_with_ids, folder_name, output_path=None):
    """
    Create a document with images arranged in a grid, 
    creating multiple pages as needed.
    
    :param image_files_with_ids: List of tuples (file_path, id)
    :param folder_name: Name of the source folder
    :param output_path: Optional path to save the output document
    """
    # Create document
    document = Document()
    
    # Modify page orientation to portrait
    sections = document.sections
    for section in sections:
        section.orientation = WD_ORIENT.PORTRAIT
        
        # Set A4 page size explicitly
        section.page_height = Cm(29.7)  # A4 height
        section.page_width = Cm(21.0)   # A4 width
    
    # Set minimal margins
    sections[0].top_margin = Cm(0.25)
    sections[0].bottom_margin = Cm(0.25)
    sections[0].left_margin = Cm(0.25)
    sections[0].right_margin = Cm(0.25)
    
    # Process images in batches of 8 (4 rows, 2 columns)
    for batch_start in range(0, len(image_files_with_ids), 8):
        # Create a new table for each page
        table = document.add_table(rows=4, cols=2)
        
        # Remove table borders by setting border width to 0
        for row in table.rows:
            for cell in row.cells:
                cell._tc.get_or_add_tcPr().append(
                    parse_xml('<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                              '<w:top w:val="nil"/>'
                              '<w:left w:val="nil"/>'
                              '<w:bottom w:val="nil"/>'
                              '<w:right w:val="nil"/>'
                              '</w:tcBorders>')
                )
        
        # Get the current batch of images (up to 8)
        batch_images = image_files_with_ids[batch_start:batch_start+8]
        
        # Insert images in the current batch
        for i, (image_path) in enumerate(batch_images):
            row = i // 2
            col = i % 2
            
            cell = table.cell(row, col)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Clear any existing paragraphs in the cell
            cell.text = ''
            
            # Add image to cell
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add picture with a consistent width that fills the page
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(4.5))
            
    
    # Generate output path if not provided
    if output_path is None:
        # Clean folder name (remove special characters and spaces)
        clean_folder_name = "".join(c if c.isalnum() else "_" for c in folder_name)
        output_path = f'IDs_{clean_folder_name}.docx'
    
    # Create output directory if it doesn't exist
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Save the document
    document.save(output_path)
    print(f"Images saved to {output_path}")

def main():
    # Prompt for image folder
    while True:
        folder_path = input("Enter the path to the folder containing images: ").strip()
        
        # Handle potential quotes around path
        folder_path = folder_path.strip("'\"")
        
        if os.path.isdir(folder_path):
            break
        else:
            print("Invalid folder path. Please try again.")
    
    
    # Get folder name from path
    folder_name = os.path.basename(os.path.normpath(folder_path))
    
    # Get image files with IDs
    image_files_with_ids = get_image_files_with_ids(folder_path)
    
    if not image_files_with_ids:
        print("No image files found in the specified folder.")
        return
    
    # Create document with images
    create_image_document(image_files_with_ids, folder_name)

if __name__ == "__main__":
    main()